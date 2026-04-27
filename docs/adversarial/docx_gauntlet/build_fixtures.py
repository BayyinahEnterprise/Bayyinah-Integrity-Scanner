"""
DOCX hidden-text adversarial gauntlet — fixture builder.

Six concealment techniques targeting the surface/substrate gap in
WordprocessingML documents. Mirrors the structure of the PDF gauntlet
(``../pdf_gauntlet/build_fixtures.py``).

The visible (surface) text is always the same Q3 financial summary
mentioning $1,000 in revenue. The hidden (substrate) text is always the
same payload mentioning $10,000. The contrast is the canonical
"contract that displays one figure and contains another" example.

Fixtures
--------

01 — White-on-white text (font color set to #FFFFFF in run-properties)
02 — Microscopic font (sz = 2 = 1.0pt; well below human-readable)
03 — Custom XML metadata (custom.xml properties carrying the payload)
04 — Comment with payload (word/comments.xml)
05 — Header/footer hidden text (separate XML parts the reader rarely sees)
06 — Footnote with payload (word/footnotes.xml)

These six are deliberately *different* from what DocxAnalyzer already
catches (vanish-flag hidden text, tracked changes, VBA, embedded files,
altChunks, external relationships, plus per-run zero-width / TAG / bidi
/ homoglyph). The point is to find the misses, not to re-prove the hits.
"""

from __future__ import annotations

import os
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "fixtures"
OUT.mkdir(parents=True, exist_ok=True)

SURFACE_TEXT = (
    "Q3 financial summary - revenue grew 8% YoY to $1,000 thousand. "
    "Margins held steady. Cash position remains strong."
)
HIDDEN_PAYLOAD = (
    "HIDDEN_TEXT_PAYLOAD: actual revenue $10,000 - see annex"
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _rezip(src_dir: Path, dest: Path) -> None:
    """Re-zip an extracted DOCX directory tree into a new .docx file.

    DOCX is a ZIP of XML parts; we mutate the parts on disk and re-pack.
    """
    if dest.exists():
        dest.unlink()
    with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, _dirs, files in os.walk(src_dir):
            for fname in files:
                full = Path(root) / fname
                arcname = full.relative_to(src_dir).as_posix()
                zf.write(full, arcname)


def _make_base_docx(path: Path) -> None:
    """Write a fresh DOCX with only the visible surface text."""
    doc = Document()
    doc.add_paragraph(SURFACE_TEXT)
    doc.save(str(path))


def _explode(docx_path: Path) -> Path:
    """Extract a DOCX into a sibling directory and return the dir path."""
    workdir = docx_path.with_suffix(".extracted")
    if workdir.exists():
        shutil.rmtree(workdir)
    workdir.mkdir()
    with zipfile.ZipFile(docx_path, "r") as zf:
        zf.extractall(workdir)
    return workdir


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def fixture_01_white_on_white() -> tuple[Path, str]:
    """Hidden text via white font color on white background.

    DocxAnalyzer does not check w:color — it only checks w:vanish.
    """
    path = OUT / "01_white_on_white.docx"
    doc = Document()
    doc.add_paragraph(SURFACE_TEXT)
    p = doc.add_paragraph()
    run = p.add_run(HIDDEN_PAYLOAD)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.save(str(path))
    return path, "White text (#FFFFFF) on default white page"


def fixture_02_microscopic_font() -> tuple[Path, str]:
    """Hidden text rendered at ~1.0pt — sub-readable.

    DocxAnalyzer does not inspect w:sz at all.
    """
    path = OUT / "02_microscopic_font.docx"
    doc = Document()
    doc.add_paragraph(SURFACE_TEXT)
    p = doc.add_paragraph()
    run = p.add_run(HIDDEN_PAYLOAD)
    run.font.size = Pt(1.0)
    doc.save(str(path))
    return path, "Run rendered at 1.0pt (well below 8pt readable threshold)"


def fixture_03_custom_xml_properties() -> tuple[Path, str]:
    """Hidden text in docProps/custom.xml.

    DocxAnalyzer never opens docProps/*. The custom-properties part is
    a documented OOXML location for arbitrary author-defined metadata
    and is read by Word and indexers but not by Bayyinah.
    """
    path = OUT / "03_custom_xml_properties.docx"
    _make_base_docx(path)
    workdir = _explode(path)

    custom_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument'
        '/2006/custom-properties" xmlns:vt="http://schemas.openxmlformats.org'
        '/officeDocument/2006/docPropsVTypes">\n'
        '  <property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" '
        'pid="2" name="ActualRevenue">\n'
        f'    <vt:lpwstr>{HIDDEN_PAYLOAD}</vt:lpwstr>\n'
        '  </property>\n'
        '</Properties>\n'
    )
    (workdir / "docProps").mkdir(exist_ok=True)
    (workdir / "docProps" / "custom.xml").write_text(custom_xml, "utf-8")

    # Register the custom part in [Content_Types].xml so Word recognises it.
    ct_path = workdir / "[Content_Types].xml"
    ct = ct_path.read_text("utf-8")
    if "custom.xml" not in ct:
        override = (
            '<Override PartName="/docProps/custom.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'custom-properties+xml"/>'
        )
        ct = ct.replace("</Types>", override + "</Types>")
        ct_path.write_text(ct, "utf-8")

    # Register the relationship in _rels/.rels (the package-level rels).
    rels_path = workdir / "_rels" / ".rels"
    rels = rels_path.read_text("utf-8")
    if "custom.xml" not in rels:
        rel = (
            '<Relationship Id="rIdCustom1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/'
            '2006/relationships/custom-properties" '
            'Target="docProps/custom.xml"/>'
        )
        rels = rels.replace("</Relationships>", rel + "</Relationships>")
        rels_path.write_text(rels, "utf-8")

    _rezip(workdir, path)
    shutil.rmtree(workdir)
    return path, "Payload in docProps/custom.xml as a custom property"


def fixture_04_comment_payload() -> tuple[Path, str]:
    """Hidden text inside a Word comment.

    DocxAnalyzer never opens word/comments.xml. Comments are visible in
    Word's review pane but disappear from the printed/exported view; an
    LLM ingesting the document may or may not consume them depending
    on the extractor.
    """
    path = OUT / "04_comment_payload.docx"
    _make_base_docx(path)
    workdir = _explode(path)

    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main">\n'
        '  <w:comment w:id="1" w:author="Reviewer" w:date="2026-04-26T00:00:00Z" '
        'w:initials="R">\n'
        f'    <w:p><w:r><w:t xml:space="preserve">{HIDDEN_PAYLOAD}</w:t>'
        '</w:r></w:p>\n'
        '  </w:comment>\n'
        '</w:comments>\n'
    )
    (workdir / "word" / "comments.xml").write_text(comments_xml, "utf-8")

    # Register comments part in [Content_Types].xml
    ct_path = workdir / "[Content_Types].xml"
    ct = ct_path.read_text("utf-8")
    if "comments.xml" not in ct:
        override = (
            '<Override PartName="/word/comments.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.comments+xml"/>'
        )
        ct = ct.replace("</Types>", override + "</Types>")
        ct_path.write_text(ct, "utf-8")

    # Register relationship in word/_rels/document.xml.rels
    drels_path = workdir / "word" / "_rels" / "document.xml.rels"
    drels = drels_path.read_text("utf-8")
    if "comments.xml" not in drels:
        rel = (
            '<Relationship Id="rIdComments1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
            'relationships/comments" Target="comments.xml"/>'
        )
        drels = drels.replace(
            "</Relationships>", rel + "</Relationships>",
        )
        drels_path.write_text(drels, "utf-8")

    _rezip(workdir, path)
    shutil.rmtree(workdir)
    return path, "Payload inside word/comments.xml (review-pane comment)"


def fixture_05_header_payload() -> tuple[Path, str]:
    """Hidden text inside a header part with white font color.

    DocxAnalyzer only parses word/document.xml. Headers and footers
    live in word/header*.xml / word/footer*.xml — a separate stream
    the analyzer never opens. We additionally render the header text
    in white so even a header-aware reader misses it visually.
    """
    path = OUT / "05_header_payload.docx"
    _make_base_docx(path)
    workdir = _explode(path)

    header_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:hdr xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main">\n'
        '  <w:p><w:r><w:rPr><w:color w:val="FFFFFF"/></w:rPr>'
        f'<w:t xml:space="preserve">{HIDDEN_PAYLOAD}</w:t></w:r></w:p>\n'
        '</w:hdr>\n'
    )
    (workdir / "word" / "header1.xml").write_text(header_xml, "utf-8")

    # Register the header part
    ct_path = workdir / "[Content_Types].xml"
    ct = ct_path.read_text("utf-8")
    if "header1.xml" not in ct:
        override = (
            '<Override PartName="/word/header1.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.header+xml"/>'
        )
        ct = ct.replace("</Types>", override + "</Types>")
        ct_path.write_text(ct, "utf-8")

    # Register the relationship from document.xml -> header
    drels_path = workdir / "word" / "_rels" / "document.xml.rels"
    drels = drels_path.read_text("utf-8")
    if "header1.xml" not in drels:
        rel = (
            '<Relationship Id="rIdHeader1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
            'relationships/header" Target="header1.xml"/>'
        )
        drels = drels.replace(
            "</Relationships>", rel + "</Relationships>",
        )
        drels_path.write_text(drels, "utf-8")

    # Wire the header into document.xml's section properties so Word
    # actually displays it. We append a referenceless headerReference
    # entry to the body's sectPr; even if Word does not pick it up, the
    # part is still in the package and any indexer that walks the ZIP
    # finds the payload.
    doc_path = workdir / "word" / "document.xml"
    doc_xml = doc_path.read_text("utf-8")
    if "<w:headerReference" not in doc_xml and "<w:sectPr" in doc_xml:
        doc_xml = doc_xml.replace(
            "<w:sectPr",
            '<w:sectPr><w:headerReference '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/'
            '2006/relationships" w:type="default" r:id="rIdHeader1"/>',
            1,
        )
        # Drop the trailing duplicate <w:sectPr> if our naive replacement
        # produced nesting; Word tolerates empty sectPr anyway.
        doc_path.write_text(doc_xml, "utf-8")

    _rezip(workdir, path)
    shutil.rmtree(workdir)
    return path, "Payload in word/header1.xml with white font color"


def fixture_06_footnote_payload() -> tuple[Path, str]:
    """Hidden text inside word/footnotes.xml.

    DocxAnalyzer never opens footnotes.xml. The fixture additionally
    omits the footnote *reference* from document.xml, so a casual
    reader of the rendered document never sees an indicator that a
    footnote exists at all — but the footnote content is in the ZIP.
    """
    path = OUT / "06_footnote_payload.docx"
    _make_base_docx(path)
    workdir = _explode(path)

    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main">\n'
        '  <w:footnote w:id="0" w:type="separator">'
        '<w:p><w:r><w:separator/></w:r></w:p></w:footnote>\n'
        '  <w:footnote w:id="1" w:type="continuationSeparator">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>\n'
        '  <w:footnote w:id="2">\n'
        f'    <w:p><w:r><w:t xml:space="preserve">{HIDDEN_PAYLOAD}</w:t>'
        '</w:r></w:p>\n'
        '  </w:footnote>\n'
        '</w:footnotes>\n'
    )
    (workdir / "word" / "footnotes.xml").write_text(footnotes_xml, "utf-8")

    ct_path = workdir / "[Content_Types].xml"
    ct = ct_path.read_text("utf-8")
    if "footnotes.xml" not in ct:
        override = (
            '<Override PartName="/word/footnotes.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.footnotes+xml"/>'
        )
        ct = ct.replace("</Types>", override + "</Types>")
        ct_path.write_text(ct, "utf-8")

    drels_path = workdir / "word" / "_rels" / "document.xml.rels"
    drels = drels_path.read_text("utf-8")
    if "footnotes.xml" not in drels:
        rel = (
            '<Relationship Id="rIdFootnotes1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
            'relationships/footnotes" Target="footnotes.xml"/>'
        )
        drels = drels.replace(
            "</Relationships>", rel + "</Relationships>",
        )
        drels_path.write_text(drels, "utf-8")

    _rezip(workdir, path)
    shutil.rmtree(workdir)
    return (
        path,
        "Payload in word/footnotes.xml with no in-body footnote reference",
    )


BUILDERS = [
    fixture_01_white_on_white,
    fixture_02_microscopic_font,
    fixture_03_custom_xml_properties,
    fixture_04_comment_payload,
    fixture_05_header_payload,
    fixture_06_footnote_payload,
]


if __name__ == "__main__":
    for builder in BUILDERS:
        path, desc = builder()
        size = path.stat().st_size
        print(f"{path.name:<40} {size:>7} bytes  - {desc}")
